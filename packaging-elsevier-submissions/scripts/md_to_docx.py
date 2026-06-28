#!/usr/bin/env python3
"""docxkit.md_to_docx —— 把投稿材料的 Markdown 转成 Editorial Manager 用的 DOCX。

由 CARETrack 的 build_docx_materials.py 泛化而来:不再写死任何论文专属路径,
改为接收任意 .md 文件列表 + 输出目录。统一 Arial 11pt 的保守排版,符合 EM 习惯。

依赖 python-docx,推荐用 uv 运行(无需预装):
    uv run --with python-docx knowledge-work/docxkit/md_to_docx.py \
        projects/<slug>/paper/submission/0_new/highlights.md \
        --out projects/<slug>/paper/submission/0_new/upload
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def iter_markdown_blocks(path: Path) -> list[tuple[str, str]]:
    """把一小撮 Markdown(标题/列表/段落)解析成带类型的块,供 DOCX 生成。"""
    blocks: list[tuple[str, str]] = []
    paragraph: list[str] = []
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            if paragraph:
                blocks.append(("p", " ".join(paragraph)))
                paragraph = []
            continue
        if line.startswith("#"):
            if paragraph:
                blocks.append(("p", " ".join(paragraph)))
                paragraph = []
            level = min(len(line) - len(line.lstrip("#")), 3)
            blocks.append((f"h{level}", line.lstrip("#").strip()))
        elif line.startswith("- ") or line.startswith("* "):
            if paragraph:
                blocks.append(("p", " ".join(paragraph)))
                paragraph = []
            blocks.append(("bullet", line[2:].strip()))
        else:
            paragraph.append(line)
    if paragraph:
        blocks.append(("p", " ".join(paragraph)))
    return blocks


def apply_document_style(document: Document) -> None:
    """套用 Arial 11pt 的保守排版,适合 EM 上传文件。"""
    styles = document.styles
    for name in ("Normal", "Body Text"):
        if name in styles:
            styles[name].font.name = "Arial"
            styles[name].font.size = Pt(11)
    for name, size in (("Title", 14), ("Heading 1", 13), ("Heading 2", 12), ("Heading 3", 11)):
        if name in styles:
            styles[name].font.name = "Arial"
            styles[name].font.size = Pt(size)


def add_block(document: Document, block_type: str, text: str) -> None:
    """把一个解析后的 Markdown 块写入 DOCX(去掉 **bold** 与 `code` 标记)。"""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    if block_type == "h1":
        document.add_heading(text, level=1).alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif block_type in ("h2", "h3"):
        document.add_heading(text, level=int(block_type[1]))
    elif block_type == "bullet":
        document.add_paragraph(text, style="List Bullet")
    else:
        document.add_paragraph(text)


def markdown_to_docx(source: Path, target: Path) -> None:
    """把单个 Markdown 文件转成一个可编辑 DOCX。"""
    document = Document()
    apply_document_style(document)
    for block_type, text in iter_markdown_blocks(source):
        add_block(document, block_type, text)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def main() -> None:
    """命令行入口:把若干 .md 转成同名 .docx 落到 --out 目录。"""
    parser = argparse.ArgumentParser(description="Markdown → EM DOCX 转换")
    parser.add_argument("sources", nargs="+", type=Path, help="一个或多个 .md 文件")
    parser.add_argument("--out", type=Path, required=True, help="输出目录")
    args = parser.parse_args()
    for src in args.sources:
        target = args.out / (src.stem + ".docx")
        markdown_to_docx(src, target)
        print(f"  {src}  ->  {target}")


if __name__ == "__main__":
    main()
